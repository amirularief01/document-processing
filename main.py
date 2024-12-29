from extract_text import extract_text_with_tika
import pandas as pd
from docx import Document
import os


def create_word(data, output_file):
    """
    Creates a Word document with a table containing the extracted data.

    :param data: List of [Header, Content] pairs to populate the table
    :param output_file: The name of the output Word document
    """
    doc = Document()

    # Add a title
    doc.add_heading('Extracted Table', level=1)

    # Add a table with two columns
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Set the headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Header'
    hdr_cells[1].text = 'Content'

    # Populate the table with data
    for header, content in data:
        row_cells = table.add_row().cells
        row_cells[0].text = header
        row_cells[1].text = content

    # Save the document
    doc.save(output_file)
    print(f"Word document created: {output_file}")



def main(input_folder, output_file):
    """
    Processes all documents in a folder, extracts text with Tika,
    and appends the data into a Word document.

    :param input_folder: Folder containing the input documents
    :param output_file: Name of the output Word document
    """
    # Collect all files in the input folder with specific extensions
    valid_extensions = ('.pdf', '.pptx', '.docx', '.txt')
    files = [f for f in os.listdir(input_folder) if f.lower().endswith(valid_extensions)]
    data = []

    # Process files one by one
    for filename in files:
        file_path = os.path.join(input_folder, filename)
        print(f"Processing file: {filename}")

        text = extract_text_with_tika(file_path)
        if not text:
            print(f"No content extracted from {filename}. Skipping...")
            continue

        # Split the extracted text into sections based on headers and content
        sections = text.split("\n\n")
        current_header = None
        current_content = []

        for section in sections:
            if section.strip().isdigit():  # Detect page numbers
                if current_header and current_content:
                    data.append([current_header, "\n".join(current_content)])
                current_header = None
                current_content = []
            elif current_header is None:  # Treat the first non-numeric section as the header
                current_header = section.strip()
            else:  # The rest is content
                current_content.append(section.strip())

        # Add the last section if any
        if current_header and current_content:
            data.append([current_header, "\n".join(current_content)])

    # Create the Word document with all data
    create_word(data, output_file)


# Specify the input folder and output file
input_folder = r"C:\Users\amiru\Documents\Self Project\notes-processing\documents"  
output_file = "Extracted_Table.docx"  # Name of the output Word file

# Run the main function
main(input_folder, output_file)
