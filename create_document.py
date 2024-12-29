from docx import Document

def create_word_document(headers_and_content, output_path):
    try:
        # Load existing document if it exists
        doc = Document(output_path)
    except:
        # Create a new document if it doesn't exist
        doc = Document()

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    for header, content in headers_and_content.items():
        row = table.add_row()
        row.cells[0].text = header
        row.cells[1].text = content

    doc.save(output_path)

