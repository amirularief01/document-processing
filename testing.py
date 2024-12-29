from extract_text import extract_text_from_pdf, extract_text_from_pptx, extract_text_with_tika
import pandas as pd
from docx import Document

#print(extract_text_with_tika(r"C:\Users\amiru\Documents\Self Project\notes-processing\documents\Integrated Quality Control 1_With Solutions(1).pptx"))
#print(extract_text_with_tika(r"C:\Users\amiru\Documents\Self Project\extraction\extract\ABF 2019 LL- (ABF - 07-E002B Leak) (1).pdf"))

text = extract_text_with_tika(r"C:\Users\amiru\Documents\Self Project\notes-processing\documents\Integrated Quality Control 1_With Solutions(1).pptx")

sections = text.split("\n\n")

data = []
current_header = None
current_content = []

# Step 2: Process each section
for section in sections:
    if section.strip().isdigit():  # Detect page numbers
        if current_header and current_content:
            # Save the current header and content
            data.append([current_header, "\n".join(current_content)])
        # Reset for the next section
        current_header = None
        current_content = []
    elif current_header is None:  # The first non-numeric section is the header
        current_header = section.strip()
    else:  # The rest is content
        current_content.append(section.strip())

# Add the last section if any
if current_header and current_content:
    data.append([current_header, "\n".join(current_content)])

# Step 3: Create a Word document
doc = Document()

# Add a title
doc.add_heading('Extracted Table', level=1)

# Add a table with two columns
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'

# Set the headers
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Header'
hdr_cells[1].text = 'Content'

# Populate the table with data
for header, content in data:
    row_cells = table.add_row().cells
    row_cells[0].text = header
    row_cells[1].text = content

# Save the document
doc.save("Extracted_Table.docx")

print("Word document created: Extracted_Table.docx")
